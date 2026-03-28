import hashlib
import json
import logging
import os
import shutil
import tempfile
from functools import lru_cache
from typing import Iterable

from django.conf import settings
from langchain_core.documents import Document
from langchain_community.vectorstores import FAISS

import docx
import pypdf

from ..config import (
    CHAT_CHUNK_OVERLAP,
    CHAT_CHUNK_SIZE,
    CHAT_EMBEDDING_DEVICE,
    CHAT_EMBEDDING_MODEL,
    CHAT_EMBEDDING_NORMALIZE,
    FAISS_ALLOW_DANGEROUS_DESERIALIZATION,
    FAISS_REQUIRE_INTEGRITY_MANIFEST,
)


logger = logging.getLogger(__name__)


def _ensure_faiss_index_dir() -> str:
    """Ensure FAISS index directory exists; migrate legacy dir if needed."""
    target_dir = str(getattr(settings, "FAISS_INDEX_DIR", ""))
    if not target_dir:
        raise RuntimeError("settings.FAISS_INDEX_DIR is not configured")

    os.makedirs(target_dir, exist_ok=True)

    # Legacy location used previously: <BASE_DIR>/ai_core/faiss_index
    base_dir = str(getattr(settings, "BASE_DIR", ""))
    legacy_dir = os.path.join(base_dir, "ai_core", "faiss_index") if base_dir else ""

    try:
        if legacy_dir and os.path.exists(legacy_dir):
            if os.path.abspath(legacy_dir) != os.path.abspath(target_dir):
                # Migrate only when target is empty (avoid clobbering existing data)
                if not os.listdir(target_dir):
                    for name in os.listdir(legacy_dir):
                        src = os.path.join(legacy_dir, name)
                        dst = os.path.join(target_dir, name)
                        try:
                            shutil.move(src, dst)
                        except Exception:
                            # Best-effort migration; continue.
                            pass
    except Exception:
        # Never break imports because of migration
        pass

    return target_dir


# Ensure dir exists at import time (keeps old behavior), with safe legacy migration.
_FAISS_DIR = _ensure_faiss_index_dir()


def _index_manifest_path(store_path: str) -> str:
    return os.path.join(store_path, "index.integrity.json")


def _sha256_file(path: str) -> str | None:
    if not os.path.exists(path):
        return None
    digest = hashlib.sha256()
    with open(path, "rb") as f:
        while True:
            chunk = f.read(1024 * 1024)
            if not chunk:
                break
            digest.update(chunk)
    return digest.hexdigest()


def _write_integrity_manifest(store_path: str) -> None:
    payload = {
        "index.faiss": _sha256_file(os.path.join(store_path, "index.faiss")),
        "index.pkl": _sha256_file(os.path.join(store_path, "index.pkl")),
    }
    with open(_index_manifest_path(store_path), "w", encoding="utf-8") as f:
        json.dump(payload, f)


def _verify_integrity_manifest(store_path: str) -> bool:
    if not FAISS_REQUIRE_INTEGRITY_MANIFEST:
        return True
    manifest_file = _index_manifest_path(store_path)
    if not os.path.exists(manifest_file):
        logger.error("FAISS integrity manifest missing: %s", manifest_file)
        return False
    try:
        with open(manifest_file, "r", encoding="utf-8") as f:
            expected = json.load(f)
    except Exception as exc:
        logger.error("Failed reading FAISS manifest: %s", exc)
        return False

    for name in ("index.faiss", "index.pkl"):
        expected_hash = expected.get(name)
        current_hash = _sha256_file(os.path.join(store_path, name))
        if expected_hash != current_hash:
            logger.error("FAISS integrity check failed for %s", name)
            return False
    return True


def _is_trusted_store_path(store_path: str) -> bool:
    try:
        base = os.path.abspath(_FAISS_DIR)
        target = os.path.abspath(store_path)
        return os.path.commonpath([base, target]) == base
    except Exception:
        return False


@lru_cache(maxsize=1)
def get_embeddings():
    """Singleton embeddings instance (avoids reloading the HF model each request)."""
    from langchain_huggingface import HuggingFaceEmbeddings

    model_name = CHAT_EMBEDDING_MODEL
    model_kwargs = {"device": CHAT_EMBEDDING_DEVICE}
    encode_kwargs = {"normalize_embeddings": CHAT_EMBEDDING_NORMALIZE}
    return HuggingFaceEmbeddings(model_name=model_name, model_kwargs=model_kwargs, encode_kwargs=encode_kwargs)


def get_faiss_cls():
    """Returns the FAISS class for vector storage."""
    return FAISS


def _get_splitter():
    from langchain_text_splitters import RecursiveCharacterTextSplitter

    return RecursiveCharacterTextSplitter(
        chunk_size=CHAT_CHUNK_SIZE,
        chunk_overlap=CHAT_CHUNK_OVERLAP,
        length_function=len,
        add_start_index=True,
    )


def _extract_documents_with_langchain(file_obj, filename: str) -> list[Document]:
    """Prefer LangChain loaders and keep a fallback path for environments missing extras."""
    lower_name = (filename or "").lower()

    suffix = ".txt"
    if lower_name.endswith(".pdf"):
        suffix = ".pdf"
    elif lower_name.endswith(".docx"):
        suffix = ".docx"

    with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp_path = tmp.name
        try:
            file_obj.seek(0)
        except Exception:
            pass
        tmp.write(file_obj.read())

    try:
        if lower_name.endswith(".pdf"):
            from langchain_community.document_loaders import PyPDFLoader

            return PyPDFLoader(tmp_path).load()

        if lower_name.endswith(".docx"):
            try:
                from langchain_community.document_loaders import Docx2txtLoader

                return Docx2txtLoader(tmp_path).load()
            except Exception:
                # Keep resilient fallback if docx2txt isn't available.
                doc = docx.Document(tmp_path)
                text = "\n".join([p.text for p in doc.paragraphs if p.text]).strip()
                if not text:
                    return []
                return [Document(page_content=text, metadata={"page": None})]

        if lower_name.endswith(".txt"):
            from langchain_community.document_loaders import TextLoader

            return TextLoader(tmp_path, encoding="utf-8", autodetect_encoding=True).load()

        return []
    finally:
        try:
            os.remove(tmp_path)
        except OSError:
            pass


def process_document_to_chunks(texts: Iterable[str], metadatas: list[dict]):
    """Split document text(s) into chunks while preserving metadata per source text."""
    splitter = _get_splitter()
    return splitter.create_documents(list(texts), metadatas=metadatas)


def build_vector_store(documents, store_path):
    """Build and save a FAISS vector store from documents."""
    vectorstore = FAISS.from_documents(documents, get_embeddings())
    vectorstore.save_local(store_path)
    _write_integrity_manifest(store_path)
    return vectorstore


_VECTOR_STORE_CACHE: dict[str, tuple[float, FAISS]] = {}


def _index_mtime(store_path: str) -> float:
    index_file = os.path.join(store_path, "index.faiss")
    if not os.path.exists(index_file):
        return 0.0
    return os.path.getmtime(index_file)


def load_vector_store(store_path: str):
    """Load an existing FAISS vector store or return None if not found."""
    if not os.path.exists(store_path) or not os.path.exists(os.path.join(store_path, "index.faiss")):
        return None
    if not _is_trusted_store_path(store_path):
        logger.error("Rejected FAISS load outside trusted dir: %s", store_path)
        return None
    if not _verify_integrity_manifest(store_path):
        logger.error("Rejected FAISS load due to integrity failure: %s", store_path)
        return None
    if not FAISS_ALLOW_DANGEROUS_DESERIALIZATION:
        logger.error("FAISS dangerous deserialization is disabled by configuration.")
        return None
    return FAISS.load_local(store_path, get_embeddings(), allow_dangerous_deserialization=True)


def get_or_load_vector_store(store_path: str):
    """Process-local cache for FAISS stores (load once per user, reload if index changes)."""
    mtime = _index_mtime(store_path)
    cached = _VECTOR_STORE_CACHE.get(store_path)
    if cached and cached[0] == mtime:
        return cached[1]
    if mtime == 0.0:
        _VECTOR_STORE_CACHE.pop(store_path, None)
        return None
    store = load_vector_store(store_path)
    if store is None:
        _VECTOR_STORE_CACHE.pop(store_path, None)
        return None
    _VECTOR_STORE_CACHE[store_path] = (mtime, store)
    return store


def invalidate_vector_store(store_path: str) -> None:
    _VECTOR_STORE_CACHE.pop(store_path, None)


def invalidate_user_docs_vector_store(user_id: int) -> None:
    store_path = os.path.join(_FAISS_DIR, f"user_docs_lc_{user_id}")
    invalidate_vector_store(store_path)


def normalize_similarity_score(score):
    """Normalize a FAISS L2 score (lower is better) to a [0, 1] similarity."""
    import math

    try:
        return math.exp(-score)
    except Exception:
        return 0.0


def process_and_index_document(user_id, file_obj, filename):
    """Extract text from PDF/DOCX/TXT, chunk it, and update the user's FAISS index."""
    lower_name = (filename or "").lower()
    page_texts: list[str] = []
    metadatas: list[dict] = []

    try:
        try:
            file_obj.seek(0)
        except Exception:
            pass

        if not lower_name.endswith((".pdf", ".docx", ".txt")):
            return {"status": "error", "message": "Unsupported file format."}

        loaded_docs = _extract_documents_with_langchain(file_obj, filename)
        if loaded_docs:
            for idx, loaded_doc in enumerate(loaded_docs):
                text = (loaded_doc.page_content or "").strip()
                if not text:
                    continue
                meta = loaded_doc.metadata or {}
                page_texts.append(text)
                metadatas.append(
                    {
                        "source": filename,
                        "original_filename": filename,
                        "page_number": meta.get("page") if meta.get("page") is not None else idx + 1,
                    }
                )

        # Fallback extraction keeps previous behavior if a loader failed silently.
        if not page_texts and lower_name.endswith(".pdf"):
            reader = pypdf.PdfReader(file_obj)
            for idx, page in enumerate(reader.pages):
                text = (page.extract_text() or "").strip()
                if not text:
                    continue
                page_texts.append(text)
                metadatas.append({"source": filename, "original_filename": filename, "page_number": idx + 1})

        if not page_texts:
            return {"status": "error", "message": "Could not extract text from file."}

        chunks = process_document_to_chunks(page_texts, metadatas)
        if not chunks:
            return {"status": "error", "message": "Text too short to index."}

        index_dir = os.path.join(_FAISS_DIR, f"user_docs_lc_{user_id}")
        vector_store = load_vector_store(index_dir)

        if vector_store:
            vector_store.add_documents(chunks)
            vector_store.save_local(index_dir)
            _write_integrity_manifest(index_dir)
        else:
            build_vector_store(chunks, index_dir)

        invalidate_user_docs_vector_store(user_id)
        return {"status": "success", "chunks_count": len(chunks)}

    except Exception as e:
        return {"status": "error", "message": str(e)}
